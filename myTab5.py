"""
LIMASCalc ist ein Programm zur Kostendarstellung und Amortisationsdarstellung eines Beleuchtungssystems
Copyright (C) [2020]  [Taras Yuzkiv, [IES] - Individual Engeneering Solutions & Adolf Schuch GmbH]
Dieses Programm ist freie Software. Sie können es unter den Bedingungen der GNU General Public License, wie von der Free Software Foundation veröffentlicht, weitergeben und/oder modifizieren, entweder gemäß Version 3 der Lizenz oder (nach Ihrer Option) jeder späteren Version.
Die Veröffentlichung dieses Programms erfolgt in der Hoffnung, daß es Ihnen von Nutzen sein wird, aber OHNE IRGENDEINE GARANTIE, sogar ohne die implizite Garantie der MARKTREIFE oder der VERWENDBARKEIT FÜR EINEN BESTIMMTEN ZWECK. Details finden Sie in der GNU General Public License.
Sie sollten ein Exemplar der GNU General Public License zusammen mit diesem Programm erhalten haben. Falls nicht, siehe <http://www.gnu.org/licenses/>.
"""
import os
import sys
import win32com.client as client

from AnyQt.QtWidgets import QFileDialog, QListWidgetItem
from PyPDF2 import PdfFileReader, PdfFileMerger
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_BREAK
from os import path

import myTab4
import LIMASCalc

class myTab5(myTab4.myTab4):
    def __init__(self):
        super().__init__()
        self.y.pushButton_3.clicked.connect(self.addDoc)
        self.y.pushButton_4.clicked.connect(self.delDoc)
        self.y.pushButton_8.clicked.connect(self.textHinzu)
        self.y.pushButton_6.clicked.connect(self.berechnungsText)
        #self.y.toolButton_13.clicked.connect(lambda: self.convertToPDF('bericht/' + str(self.b.projektname) + '-Bericht'))
        self.y.pushButton_5.clicked.connect(
            lambda: self.testConv(os.path.dirname(sys.argv[0]) + "/Bericht/" + str(self.b.projektname) + '-Bericht'))
        self.y.toolButton_13.clicked.connect(
            lambda: self.makeMyPDF(os.path.dirname(sys.argv[0]) + "/Bericht/" + str(self.b.projektname) + '-allesZusammen.pdf'))
        #self.y.toolButton_13.clicked.connect(self.createDoc)

    def makeMyPDF(self, zielAdr):
        itemsTextList = [str(self.listWidget.item(i).text()) for i in range(self.listWidget.count())]
        mypdf = PdfFileMerger()
        for k in itemsTextList:
            mypdf.append(PdfFileReader(str(k)), 'rb')
        mypdf.write(zielAdr)
        os.startfile(zielAdr)


    def testConv(self, meineAdr):
        """Save a pdf of a docx file."""
        try:
            meineDocx = meineAdr + '.docx'
            word = client.DispatchEx("Word.Application")
            target_path = meineDocx.replace(".docx", r".pdf")
            if path.exists(target_path):
                buttonReply = LIMASCalc.QMessageBox.question(self, 'Achtung',
                                                   "Bericht existiert bereits, soll die PDF ersetzt werden?",
                                                   LIMASCalc.QMessageBox.Yes | LIMASCalc.QMessageBox.No, LIMASCalc.QMessageBox.No)
                if buttonReply == LIMASCalc.QMessageBox.Yes:
                    os.remove(target_path)
                    self.infoPopUp("Die Datei wurde erfolgreich ersetzt.", "Erledigt")
                    word_doc = word.Documents.Open(meineDocx)
                    word_doc.SaveAs(target_path, FileFormat=17)
                    word_doc.Close()
                    #os.startfile(target_path)
                    word.Quit()
                else:
                    pass
            else:
                word_doc = word.Documents.Open(meineDocx)
                word_doc.SaveAs(target_path, FileFormat=17)
                word_doc.Close()
                #os.startfile(target_path)
                word.Quit()
            #self.makeMyPDF(target_path)
            item = QListWidgetItem(target_path)
            self.y.listWidget.addItem(item)
            self.y.listWidget.scrollToBottom()
        except Exception as e:
            raise e


    def addDoc(self):
        self.y.lineEdit_28.clear()
        fname, _ = QFileDialog.getOpenFileName(self, 'Open file', 'c:\\', "Image files (*.pdf)")
        self.fname = fname
        self.y.lineEdit_28.insert(os.path.realpath(self.fname))

    def delDoc(self):
        self.listWidget.takeItem(self.y.listWidget.currentRow())

    def textHinzu(self):
        self.meinText = self.y.lineEdit_28.text()
        item = QListWidgetItem(self.meinText)
        self.y.listWidget.addItem(item)
        self.y.listWidget.scrollToBottom()

    def openWord(self):
        temp = str(os.path.dirname(sys.argv[0]))
        lauf = str(temp+'/bericht/' + str(self.b.projektname) + '-Bericht.docx')
        os.startfile(lauf)

    def ohneLIMASmitLED(self):
        document = Document('bericht/berichtVorlage.docx')
        table = document.add_table(rows=4, cols=4)
        #table.style = 'Colorful Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].paragraphs[0].add_run('Allgemein').bold = True
        hdr_cells[2].paragraphs[0].add_run('Ansprechpartner').bold = True
        hdr_cells = table.rows[1].cells
        hdr_cells[0].text = ('Firmenname')
        hdr_cells[1].text = self.b.firmenname
        hdr_cells[2].text = 'Name'
        hdr_cells[3].text = self.b.name

        run = document.add_paragraph().add_run()
        font = run.font
        font.name = 'Calibri'
        font.size = Pt(20)
        cells1 = table.rows[1].cells
        cells1[0].text = self.b.firmenname
        document.save('bericht/' + str(self.b.projektname) + '-Bericht.docx')



    def berechnungsText(self):
        abblindungnr = 0
        document = Document('bericht/berichtVorlage.docx')

        paragraph01 = document.add_paragraph()
        paragraph01.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        run04 = paragraph01.add_run("")
        font01 = run04.font
        font01.name = 'Calibri'
        font01.size = Pt(10)
        run04.add_text(str(self.b.zeitDatum))
        run04.add_break()

        table = document.add_table(rows=4, cols=2)
        table.style.font.name = 'Calibri'
        table.style = 'Table Grid'

        hdr_cells = table.rows[0].cells
        hdr_cells[0].paragraphs[0].add_run('Allgemein').bold = True
        hdr_cells[1].paragraphs[0].add_run('Ansprechpartner').bold = True

        hdr_cells = table.rows[1].cells
        hdr_cells[0].text = ('Firmenname: '+ str(self.b.firmenname))
        hdr_cells[1].text = "Name: " + str(self.b.name)

        hdr_cells = table.rows[2].cells
        hdr_cells[0].text = ('Projektname: ' + str(self.b.projektname))
        hdr_cells[1].text = "Email: " + str(self.b.email)

        hdr_cells = table.rows[3].cells
        hdr_cells[0].text = ('Strompreis in € pro kWh: ' + str(self.b.strompreis))
        hdr_cells[1].text = "Telefon: " + str(self.b.tel)

        p1 = document.add_paragraph()
        p1.style.font.name = 'Calibri'
        p1.style.font.size = Pt(10)
        run1 = p1.add_run()
        run1.add_break()
        run1.add_break()

        if (self.b.bewegungsmelder != False) or (self.b.tageslicht != False) or (self.b.kalenderCheck != False):
            p1.add_run("Eine Kosten- und Amortisationsberechnung der Beleuchtungsanlage mit Schuch Leuchten und einem Lichtmanagentsystem.").bold = True
        else:
            p1.add_run("Eine Kosten- und Amortisationsberechnung der Beleuchtungsanlage mit Schuch Leuchten.").bold = True

        p2 = document.add_paragraph()
        run2 = p2.add_run()
        run2.add_break()
        p2.add_run("Die neue Beleuchtungsanlage der Firma Adolf Schuch GmbH wird mit insgesamt " + str(self.b.neuAnzahl) + " Leuchten des Typs " + str(self.b.neuModell) + " betrachtet. Eine Leuchte weist eine Leistung von " + str(self.b.neuLeistung) + "Watt auf. "
                    "Pro Jahr entsteht dabei ein Gesamtverbrauch von " + str(int(self.b.verbrauchNeu)) + "kWh. Dabei wird von einer Betriebsdauer von " + str(self.b.neuhProT) + "h pro Tag ausgegangen. An " + str(self.b.neutProJ) +
                   " der 365 Tagen eines Jahres wird die Beleuchtung betrieben. ")
        run21 = p2.add_run()
        run21.add_break()
        font010 = run2.font
        font010.name = 'Calibri'
        font010.size = Pt(10)

        paragraph = document.add_paragraph()
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run3 = paragraph.add_run("")
        font = run3.font
        font.name = 'Calibri'
        font.size = Pt(10)
        run3.add_picture('pics/leuchte.jpg', height=Inches(1.5))
        run3.add_break()
        abblindungnr+=1
        run3.add_text("Abbildung " + str(abblindungnr) + ": Schuch Leuchte, " + str(self.b.neuModell))
        run31 = paragraph.add_run()
        run31.add_break()

        p3 = document.add_paragraph()
        p3.add_run("Die Bestandanlage mit insgesamt " + str(self.b.altAnzahl) + " Leuchten des Typs " + str(self.b.altModell) + " der Firma " + str(self.b.altHersteller) + " weist einen jährlichen Gesamtverbrauch von " + str(self.b.verbrauchAlt)
                          + "kWh auf. Eine Leuchte hat dabei eine Leistung von " + str(self.b.altLeistung) + "Watt.")
        p3.add_run(" Auf dem folgenden Diagram ist eine Gegenüberstellung des Verbrauchs der Bestandsanlage (dargestellt mit rot) und der Beleuchtung mit Schuch Leuchten (dargestellt mit blau) zu sehen. Durch die Bestandsanlage entstehen pro Jahr, bei einem Strompreis"
                   "von "+ str(self.b.strompreis) +"Cent/kWh, insgesamt " + "Kosten von "+ str(round(self.b.stromUndReparaturAlt, 2)) +" Euro. Jährliche Reparaturkosten mit inbegriffen. Duruch die Schuch Leuchten können die Gesamtkosten auf " + str(round(self.b.stromUndReparaturNeu, 2)) + "€ gesenkt werden. Dabei entsteht eine jährliche Ersparnis"+
                    " von " + str(round(self.b.ersparnisAltNeuMinusZusatzkosten, 2)) + "Euro bzw. eine Ersparnis von " + str(round(100-(100*round(self.b.stromUndReparaturNeu)/round(self.b.stromUndReparaturAlt)), 1)) + "%. Anfallende Zusatzkosten inklusive. ")

        if ((self.b.bewegungsmelder != False) or (self.b.tageslicht != False) or (self.b.kalenderCheck != False)) and self.y.checkBox_9.isChecked():
            p3.add_run().add_break()
            p3.add_run().add_break()
            p3.add_run("Durch das intelligente Lichtmanagementsystem (LIMAS) werden die Stromkosten weiter reduzuert. Mit Schuch Leuchten und einem Lichtmanagementsystem werden jährlich insgesamt " + str(round(self.b.GesamtverbrauchInkW)) +
                       "kW verbraucht. Dadurch reduzieren sich die jährlichen Stromkosten auf " + str(round(self.b.Gesamtkosten)) + "€. Dies ergibt, vergliechen mit der Anlage ohne Lichtmanagement, eine Kostenersparnis (eventuelle jährliche Zusatzkosten mit inbegriffen) von "
                       + str(round(self.b.minusJährlicheZusatzkosten)) + "€. Verglichen mit der Beleuchtungsanlage ohne ein Lichtmanagementsystem ist bei dieser Konfiguration eine Ersparnis bis zu "+ str(round((100*self.b.minusJährlicheZusatzkosten)/(self.b.minusJährlicheZusatzkosten+self.b.Gesamtkosten), 1)) + "% erzielt werden.")

        paragraph1 = document.add_paragraph()
        paragraph1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run4 = paragraph1.add_run("")
        font1 = run4.font
        font1.name = 'Calibri'
        font1.size = Pt(10)
        run4.add_picture('pics/bild1Tab2.jpg', height=Inches(3))
        run4.add_break()
        abblindungnr += 1
        run4.add_text("Abbildung " + str(abblindungnr) + ": Gegenüberstellung der Kosten Bestandsanlage und Schuch Leuchten")
        run41 = paragraph1.add_run()
        run41.add_break()

        p4 = document.add_paragraph()
        run42 = p4.add_run("Abbildung 3 zeigt auf, wieviel Ersparnis die neue Beleuchtung gegenüber der Bestandsanlage aufweist. ")
        run42.add_break()

        paragraph2 = document.add_paragraph()
        paragraph2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run5 = paragraph2.add_run("")
        font1 = run5.font
        font1.name = 'Calibri'
        font1.size = Pt(10)
        run5.add_picture('pics/bild1Tab.jpg', height=Inches(3))
        run5.add_break()
        abblindungnr += 1
        run5.add_text("Abbildung " + str(abblindungnr) + ": Ersparnis gegenüber der Bestandsanlage")
        run51 = paragraph2.add_run()
        run51.add_break()


        if self.y.checkBox_4.isChecked():
            paragraph004 = document.add_paragraph()
            paragraph004.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run004 = paragraph004.add_run("")
            font004 = run004.font
            font004.name = 'Calibri'
            font004.size = Pt(14)
            run004.bold = True
            run004.italic = True
            run004.underline = True
            run004.add_break()
            run004.add_break(WD_BREAK.PAGE)
            if ((self.b.bewegungsmelder != False) or (self.b.tageslicht != False) or (self.b.kalenderCheck != False)):
                run004.add_text(
                    "Amortisationsberechnung der Schuch Leuchten mit einem Lichtmanagemetsystem im Vergleich mit der Bestandanlage")
                paragraph005 = document.add_paragraph()
                paragraph005.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                run005 = paragraph005.add_run("")
                font005 = run005.font
                font005.name = 'Calibri'
                font005.size = Pt(10)
                run005.add_break()
                run005.add_text(
                    "Abbildung " + str(abblindungnr+1) + " stellt die Amortisationszeiträume der Anlage gegenüber. Der rote Balken stellt den Amortisationszeitraum der Beleleuchtungsanlge mit Schuch Leuchten dar. Der blaue Balken "
                                                         "gibt den Amortisationszeitraum der Anlage mit Schuch Leuchten und einem Lichtmanagementsystem wieder. Der grüne Balken beschreibt die Amortisationszeitdauer des Lichtmanagementsystems"
                                                         " ausschließlich durch das Lichtmanagementsystem. Hierbei wird durch die Schuch Leuchten verursachte Kostenersparnis nicht miteinbezogen. Es werden die Anschaffungskosten des gesamten Lichtmanagementsystems, "
                                                         "samt dem Leuchtenaufpreis für Lichtmanagentsystemfähige Leuchten, mit den eingesparten Kosten durch das Lichtmanagementsystem verrechnet. Alle Angaben erfolgen in Jahren. ")
            else:
                run004.add_text(
                    "Amortisationsberechnung von Schuch Leuchten im Vergleich mit der Bestandsanlage")
                paragraph005 = document.add_paragraph()
                paragraph005.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                run005 = paragraph005.add_run("")
                font005 = run005.font
                font005.name = 'Calibri'
                font005.size = Pt(10)
                run005.add_break()
                run005.add_text("Abbildung " + str(abblindungnr+1) + " zeigt den Amortissationszeitraum der Beleuchtungsanlage mit Schuch Leuchten. Durch eine intelligente Lichtsteuerung/Lichtregelung "
                                                                     "können Stromkosten weiter reduzuert werden. Alle Angaben erfolgen in Jahren.")
            paragraph3 = document.add_paragraph()
            paragraph3.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run6 = paragraph3.add_run("")
            font2 = run6.font
            font2.name = 'Calibri'
            font2.size = Pt(10)
            run6.add_picture('pics/bild1Tab3.jpg', height=Inches(3))
            run6.add_break()
            abblindungnr += 1
            run6.add_text("Abbildung " + str(abblindungnr) + ": Amortisationszeiträume in Jahren")
            run61 = paragraph3.add_run()
            run61.add_break()
            run61.add_break()

        if self.y.checkBox_12.isChecked():
            paragraphScreen = document.add_paragraph()
            paragraphScreen.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            screenShot = paragraphScreen.add_run("")
            font0021 = screenShot.font
            font0021.name = 'Calibri'
            font0021.size = Pt(14)
            screenShot.add_break(WD_BREAK.PAGE)
            screenShot.add_text(
                "Bei der Berechnung verwendeten Parameter")
            screenShot.add_break()
            screenShot1 = document.add_paragraph()
            screenShot1.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            screenShot1 = screenShot1.add_run("")
            font0031 = screenShot1.font
            font0031.name = 'Calibri'
            font0031.size = Pt(10)
            screenShot1.add_text(
                "Im Folgenden sind die im Berechnungstool eingestellten Parameter bzw. eingegebene Werte. Einige dieser Angaben basieren auf Erfahrungswerten. Je genauer die Gegebenheiten bekannt sind, umso genauer kann die Kosten-/Amortisationsrechnung durchgeführt werden.")
            screenShot1.add_break()
            screenShot.bold = True
            screenShot.italic = True
            screenShot.underline = True
            paragraph123 = document.add_paragraph()
            paragraph123.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run321 = paragraph123.add_run("")
            font321 = run321.font
            font321.name = 'Calibri'
            font321.size = Pt(10)
            run321.add_picture('pics/shot_tab2.jpg', width=Inches(6.5))
            run321.add_break()
            abblindungnr += 1
            run321.add_text("Abbildung " + str(abblindungnr) + ": Angaben zu den Leuchten")
            run321.add_break()
            run321 = paragraph123.add_run()
            run321.add_break()
            run321.add_picture('pics/shot_tab3.jpg', width=Inches(6.5))
            abblindungnr += 1
            run321.add_text("Abbildung " + str(abblindungnr) + ": Angaben zum Lichtmanagement")
            #run321.add_break(WD_BREAK.PAGE)

        if self.y.checkBox_8.isChecked():
            paragraph002 = document.add_paragraph()
            paragraph002.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run002 = paragraph002.add_run("")
            font002 = run002.font
            font002.name = 'Calibri'
            font002.size = Pt(14)
            run002.add_break(WD_BREAK.PAGE)
            if (self.b.bewegungsmelder != False) or (self.b.tageslicht != False) or (self.b.kalenderCheck != False):
                run002.add_text(
                    "Emissionsreduzierung durch Schuch Leuchten und das Lichtmanagementsystem")
                paragraph003 = document.add_paragraph()
                paragraph003.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                run003 = paragraph003.add_run("")
                font003 = run003.font
                font003.name = 'Calibri'
                font003.size = Pt(10)
                run003.add_text(
                    "Im Folgenden sind die Emissionswerte der Bestandsanlage (in rot), der Anlage mit Schuch Leuchten (in gelb) und Schuch Leuchten mit Lichtmanagementsystem (in grün) dargestellt. Alle Angeben (außer Quecksilber) erfolgen in Kilogram pro Jahr. Quecksilber wird "
                    "in Gramm pro Jahr ausgewiesen.")
            else:
                run002.add_text(
                    "Emissionsreduzierung durch Schuch Leuchten")
                paragraph003 = document.add_paragraph()
                paragraph003.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                run003 = paragraph003.add_run("")
                font003 = run003.font
                font003.name = 'Calibri'
                font003.size = Pt(10)
                run003.add_text(
                    "Im Folgenden sind die Emissionswerte der Bestandsanlage (in rot) sowie der Anlage mit Schuch Leuchten (in gelb) dargestellt. Alle Angeben (außer Quecksilber) erfolgen in Kilogram pro Jahr. Quecksilber wird "
                    "in Gramm pro Jahr ausgewiesen.")
            run002.bold = True
            run002.italic = True
            run002.underline = True
            run002.add_break()


            paragraph4 = document.add_paragraph()
            paragraph4.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run6 = paragraph4.add_run("")
            font3 = run6.font
            font3.name = 'Calibri'
            font3.size = Pt(10)
            run6.add_picture('pics/bild1Tab3_Kohlendioxid.jpg', height=Inches(3))
            run6.add_break()
            abblindungnr += 1
            run6.add_text("Abbildung " + str(abblindungnr) + ": Vergleich Ausstoß Kohlendioxid")
            run62 = paragraph4.add_run()
            run62.add_break()

            paragraph5 = document.add_paragraph()
            paragraph5.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run7 = paragraph5.add_run("")
            font4 = run7.font
            font4.name = 'Calibri'
            font4.size = Pt(10)
            run7.add_picture('pics/bild1Tab3_Kohlenmonoxid.jpg', height=Inches(3))
            run7.add_break()
            abblindungnr += 1
            run7.add_text("Abbildung " + str(abblindungnr) + ": Vergleich Ausstoß Kohlenmonoxid")
            run71 = paragraph5.add_run()
            run71.add_break()

            paragraph6 = document.add_paragraph()
            paragraph6.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run8 = paragraph6.add_run("")
            font5 = run8.font
            font5.name = 'Calibri'
            font5.size = Pt(10)
            run8.add_picture('pics/bild1Tab3_Lachgas.jpg', height=Inches(3))
            run8.add_break()
            abblindungnr += 1
            run8.add_text("Abbildung " + str(abblindungnr) + ": Vergleich Ausstoß Lachgas")
            run81 = paragraph6.add_run()
            run81.add_break()

            paragraph7 = document.add_paragraph()
            paragraph7.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run9 = paragraph7.add_run("")
            font6 = run9.font
            font6.name = 'Calibri'
            font6.size = Pt(10)
            run9.add_picture('pics/bild1Tab3_Methan.jpg', height=Inches(3))
            run9.add_break()
            abblindungnr += 1
            run9.add_text("Abbildung " + str(abblindungnr) + ": Vergleich Ausstoß Methan")
            run91 = paragraph7.add_run()
            run91.add_break()

            paragraph8 = document.add_paragraph()
            paragraph8.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run10 = paragraph8.add_run("")
            font7 = run10.font
            font7.name = 'Calibri'
            font7.size = Pt(10)
            run10.add_picture('pics/bild1Tab3_NMVOC.jpg', height=Inches(3))
            run10.add_break()
            abblindungnr += 1
            run10.add_text("Abbildung " + str(abblindungnr) + ": Vergleich Ausstoß NMVOC")
            run101 = paragraph8.add_run()
            run101.add_break()

            paragraph9 = document.add_paragraph()
            paragraph9.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run11 = paragraph9.add_run("")
            font8 = run11.font
            font8.name = 'Calibri'
            font8.size = Pt(10)
            run11.add_picture('pics/bild1Tab3_PM10.jpg', height=Inches(3))
            run11.add_break()
            abblindungnr += 1
            run11.add_text("Abbildung " + str(abblindungnr) + ": Vergleich Ausstoß PM10")
            run111 = paragraph9.add_run()
            run111.add_break()

            paragraph10 = document.add_paragraph()
            paragraph10.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run12 = paragraph10.add_run("")
            font9 = run12.font
            font9.name = 'Calibri'
            font9.size = Pt(10)
            run12.add_picture('pics/bild1Tab3_Quecksilber.jpg', height=Inches(3))
            run12.add_break()
            abblindungnr += 1
            run12.add_text("Abbildung " + str(abblindungnr) + ": Vergleich Ausstoß Quecksilber")
            run121 = paragraph10.add_run()
            run121.add_break()

            paragraph11 = document.add_paragraph()
            paragraph11.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run13 = paragraph11.add_run("")
            font10 = run13.font
            font10.name = 'Calibri'
            font10.size = Pt(10)
            run13.add_picture('pics/bild1Tab3_Schwefeldioxid.jpg', height=Inches(3))
            run13.add_break()
            abblindungnr += 1
            run13.add_text("Abbildung " + str(abblindungnr) + ": Vergleich Ausstoß Schwefeldioxid")
            run131 = paragraph11.add_run()
            run131.add_break()

            paragraph12 = document.add_paragraph()
            paragraph12.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run14 = paragraph12.add_run("")
            font11 = run14.font
            font11.name = 'Calibri'
            font11.size = Pt(10)
            run14.add_picture('pics/bild1Tab3_Staub.jpg', height=Inches(3))
            run14.add_break()
            abblindungnr += 1
            run14.add_text("Abbildung " + str(abblindungnr) + ": Vergleich Ausstoß Staub")
            run141 = paragraph12.add_run()
            run141.add_break()

            paragraph13 = document.add_paragraph()
            paragraph13.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run15 = paragraph13.add_run("")
            font12 = run15.font
            font12.name = 'Calibri'
            font12.size = Pt(10)
            run15.add_picture('pics/bild1Tab3_Stickstoffdioxid.jpg', height=Inches(3))
            run15.add_break()
            abblindungnr += 1
            run15.add_text("Abbildung " + str(abblindungnr) + ": Vergleich Ausstoß Stickstoffdioxid")
            run151 = paragraph13.add_run()
            run151.add_break()
            run151.add_break()
            run151.add_break()

        if self.y.checkBox_6.isChecked():
            paragraph001 = document.add_paragraph()
            paragraph001.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run001 = paragraph001.add_run("")
            font001 = run001.font
            font001.name = 'Calibri'
            font001.size = Pt(14)
            run001.add_break(WD_BREAK.PAGE)
            run001.add_text(
                "Zu den Emissionsfaktoren")
            run001.bold = True
            run001.italic = True
            run001.underline = True

            paragraph15 = document.add_paragraph()
            paragraph15.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run17 = paragraph15.add_run("")
            font14 = run17.font
            font14.name = 'Calibri'
            font14.size = Pt(10)
            run17.add_picture('pics/2020-04-01_uba_tabelle_spezifische_ef_strommix_luftschadstoffe_update_2018.png', height=Inches(3))
            run17.add_break()
            abblindungnr += 1
            run17.add_text("Abbildung " + str(abblindungnr) + ": Spezifische Emissionsfaktoren für den deutschen Strommix")
            #run171 = paragraph15.add_run()
            #run171.add_break()

            p111 = document.add_paragraph()
            p111.style.font.name = 'Calibri'
            run111 = p111.add_run()

            p111.add_run("Die Abbildung " + str(abblindungnr) + " zeigt für die Jahre 1990, 2000 und für das Jahr 2017 die Emissionen pro Kilowattstunde Strom (auch spezifische Emissionen genannt) des deutschen Strommixes für ausgewählte Schadstoffe. Diese Werte enthalten nur die direkten Emissionen der Stromerzeugung. Durch Vorketten, wie zum Beispiel bei der Förderung der Brennstoffe, entstandene Emissionen werden in dieser Auswertung nicht berücksichtigt. Dies entspricht den internationalen Vorgaben für die Emissionsberichterstattung der Treibhausgase und der Luftschadstoffe. "
                           "Insgesamt wurden die Emissionen der meisten Schadstoffe aus der Stromerzeugung im Vergleich zu 1990 deutlich reduziert. Die stärksten Minderungen erfolgten in den neuen Bundesländern und wirkten sich vor allem bei Staub, Schwefeldioxid und Stickstoffdioxid aus. Auch die Kohlenmonoxid- und die Quecksilberemissionen sanken im betrachteten Zeitraum. In der damaligen DDR verfügten die Stromerzeugungsanlagen über keine oder allenfalls nur eine rudimentäre Abgasreinigung. Dadurch waren die Schadstoffemissionen bis zum Jahr 1990 sehr hoch. "
                           "Im Zuge der Wiedervereinigung kam es zur Schließung einer Reihe von Altanlagen. Ein kleiner Teil der Kapazitäten wurde durch Neuanlagen ersetzt, die über den neuesten Stand der Technik verfügten. Einige Altanlagen wurden technisch nachgerüstet, um die gesetzlichen Anforderungen an die Luftreinhaltung zu erfüllen. In den alten Bundesländern verfügten die großen Stromerzeugungsanlagen bereits vor 1990 über entsprechende Abgasreinigungssysteme. "
                           "Allerdings kam es auch hier aufgrund von Grenzwertverschärfungen zu einer Verringerung der Emissionswerte. Daher sinken die spezifischen Emissionen für einen großen Teil der Luftschadstoffe."
                           "Lediglich die flüchtigen organischen Verbindungen ohne Methan (NMVOC) und vor allem die Methanemissionen steigen demgegenüber seit 1990 an. Ein wesentlicher Faktor ist der seit 2004 beobachtete massive Zubau an Biogasanlagen zur Stromerzeugung. Biogas wird im Wesentlichen in Motoren verbrannt. Dadurch entstehen höhere NMVOC- und Methanemissionen als bei anderen Verbrennungsarten. Für das Treibhausgas Lachgas kann aufgrund der geringen Anzahl der Messwerte kein eindeutiger Trend bestimmt werden. "
                           "Es wird davon ausgegangen, dass sich die spezifischen Emissionen nicht verändert haben. "
                           "Die Feinstaub emissionen werden für das Jahr 1990 nicht berichtet, da keine belastbaren Informationen zur Staubzusammensetzung in den frühen 1990er Jahren vorliegen. Es ist sehr wahrscheinlich, dass diese sich von den heutigen Werten unterscheiden. Aus diesem Grund kann kein Trend bestimmt werden. Die Entwicklung zwischen dem Jahr 2000 und dem aktuellen Berichtsjahr orientiert sich an der Entwicklung der Staubemissionen, die seit dem leicht gesunken sind.")#.bold = True
            run111.add_break()
            paragraph16 = document.add_paragraph()
            paragraph16.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            run18 = paragraph16.add_run("")
            font15 = run18.font
            font15.name = 'Calibri'
            font15.size = Pt(10)
            paragraph16.add_run("Quelle: Umweltbundesamt").italic = True

            paragraph0001 = document.add_paragraph()
            paragraph0001.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run0001 = paragraph0001.add_run("")
            font0001 = run0001.font
            font0001.name = 'Calibri'
            font0001.size = Pt(14)
            run0001.add_break(WD_BREAK.PAGE)
            run0001.add_text(
                "Bemerkungen zu den Ergebnissen")
            run0001.bold = True
            run0001.italic = True
            run0001.underline = True


            p1111 = document.add_paragraph()
            p1111.style.font.name = 'Calibri'
            run1111 = p1111.add_run()

            p1111.add_run("Alle Angaben basieren auf theoretisch berechneten Werten. Adolf Schuch GmbH übernimmt keine Verantwortung für die Richtigkeit der berechneten und angezeigten Ergebnisse. Es kann nicht garantiert werden, "
                          "dass der berechnete Verbrauch bzw. berechnete Kosten den tatsächlichen Kosten entsprechen. Bei diesen Berechnungen wurde die Überschreitung der gemessenen Leistung einer Leuchte, welche nach der DIN EN 62722-1 bei bis zu 10% "
                          "liegen kann, nicht berücksichtigt. Des Weiteren, würde bei den Berechnung angenommen, dass die Leistung der Leuchte sich proportional zum Dimmwert der Leuchte verhält. Der Wirkungsgrad der EVGs ist in den meisten Fällen beim "
                          "maximalen Strom, am höchsten. Beim Verringern des Stromes (Dimmen der Leuchte) verringert sich der Wikrungsgrad ebenso. Dies hat zufolge, dass bei einem Dimmwert von zum Beispiel 10% nicht 10% der Leistung gemessen werden. \n\n"
                          "Hinsichtlich des Lichtmanagementsystems können ebenso lediglich Annahmen getroffen werden. So kann zum Beispiel die Frequentierung eines Raumes nicht genau angegeben werden. Der Tageslichtanteil in einem Raum hängt von vielen "
                          "Faktoren ab und ist jahreszeit- und wetterabhängig."
                          "")
            run1111.add_break()


        try:
            if os.path.exists(
                    os.path.dirname(sys.argv[0]) + '/bericht/' + str(self.b.projektname) + '-Bericht.docx'):
                buttonReply = LIMASCalc.QMessageBox.question(self, 'Achtung',
                                                             "Bericht existiert bereits, soll die .docx ersetzt werden?",
                                                             LIMASCalc.QMessageBox.Yes | LIMASCalc.QMessageBox.No,
                                                             LIMASCalc.QMessageBox.No)
                if buttonReply == LIMASCalc.QMessageBox.Yes:
                    os.remove(
                        os.path.dirname(sys.argv[0]) + '/bericht/' + str(self.b.projektname) + '-Bericht.docx')
                    document.save('bericht/' + str(self.b.projektname) + '-Bericht.docx')
                    self.infoPopUp("Die Datei wurde erfolgreich ersetzt.", "Erledigt")
                    self.openWord()
                else:
                    pass
            else:
                document.save('bericht/' + str(self.b.projektname) + '-Bericht.docx')
                self.openWord()
        except:
            self.infoPopUp("Bitte erst die Word Datei schließen, danach erneut versuchen zu öffnen.", "Datei kann nicht geöffnet werden")